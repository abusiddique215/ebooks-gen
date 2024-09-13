import json
from docx import Document

# Load the JSON data
with open('ebook_content.json', 'r') as file:
    ebook_data = json.load(file)

# Create a new Word document
doc = Document()

# Add the book title
doc.add_heading(ebook_data['title'], 0)

# Iterate through chapters and subsections
for chapter in ebook_data['chapters']:
    doc.add_heading(chapter['title'], level=1)
    
    for subsection in chapter['subsections']:
        doc.add_heading(subsection['title'], level=2)
        
        # Split content into paragraphs and add them
        paragraphs = subsection['content'].split('\n')
        for para in paragraphs:
            p = doc.add_paragraph()
            # Split paragraph into runs (for potential formatting)
            words = para.split()
            for word in words:
                if word.startswith('*') and word.endswith('*'):
                    p.add_run(word.strip('*')).italic = True
                elif word.startswith('**') and word.endswith('**'):
                    p.add_run(word.strip('**')).bold = True
                else:
                    p.add_run(word + ' ')

    # Add a page break after each chapter
    doc.add_page_break()

# Save the Word document
doc.save('ebook.docx')

print("Word document conversion complete. File saved as ebook.docx")